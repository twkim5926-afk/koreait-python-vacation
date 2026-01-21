import requests
from bs4 import BeautifulSoup
import time
import os
from docx import Document

# https://de.mofa.go.kr/de-ko/brd/m_7204/view.do?seq=2975117

BASE_ID = 2975117
BASE_URL = "https://de.mofa.go.kr/de-ko/brd/m_7204/view.do?seq="
count = 10
target_urls = []
for num in range(0,10):
    print(f"{BASE_URL}{BASE_ID - num}")
    target_urls.append(f"{BASE_URL}{BASE_ID - num}")

for url in target_urls:
    response = requests.get(url)

    # 통신성공 코드
    # 200이면 통신은 성공한것
    print(response.status_code)

    soup = BeautifulSoup(response.text, "html.parser")

    # 게시글 제목, 작성자, 작성일
    # select는 언제나 리스트로 리턴
    # 하나만 가져오고 싶을때 select_one
    # 위에서 첫번째 발견된 bo_head 클래스를 가진 태그
    header_tag = soup.select_one(".bo_head")

    # 게시글 누락 가능성존재(패턴에서 벗어나는 경우)
    # 기존 페이지가 아닌 에러페이지라면
    if header_tag is None:
        print(" -> 에러페이지 발생!! 건너뜁니다")
        continue

    # 타이틀 추출
    title_tag = header_tag.select_one("h2")
    title = title_tag.text.strip()
    print(f"제목:{title}")

    # 작성자, 작성일 -> dl 태그 안에 dd 태그들
    dd_tags = header_tag.select("dl dd")
    # 저자 추출
    author_tag = dd_tags[0]
    author = author_tag.text.strip()
    print(f"저자:{author}")
    # 작성일 추출
    date_tag = dd_tags[1]
    date = date_tag.text.strip()
    print(f"작성일:{date}")




    # 상위태그의.text에는 하위태그의 text들도 모두 포함
    # 도전! - 본문내용 추출
    contents_tag = soup.select_one(".bo_con")
    contents = contents_tag.text.strip()
    print(f"내용: {contents[:100]}...")

    ##### docx 파일로 저장 #####
    ROOT_PATH = "./독일대사관_게시글" # 저장할 폴더위치
    doc = Document()

    # 인코딩(폰트) 전용 함수
    def apply_font(*texts):
        for text in texts:
            if text.runs:
                text.runs[0].font.name = "Malgun Gothic"

    doc_title = doc.add_heading(title, level=1)
    pg1 = doc.add_paragraph(f"작성자: {author}")
    pg2 = doc.add_paragraph(f"작성일자: {date}")
    doc.add_paragraph("")
    pg3 = doc.add_paragraph(contents)

    apply_font(doc_title, pg1, pg2, pg3)

    # 폴더 생성
    os.makedirs(ROOT_PATH, exist_ok=True)

    # 파일경로 조립
    file_path = f"{ROOT_PATH}/{title}.docx"

    # 저장
    doc.save(file_path)
    print(f" -> 저장완료: {file_path}")
    time.sleep(1.5)




    time.sleep(1.5)